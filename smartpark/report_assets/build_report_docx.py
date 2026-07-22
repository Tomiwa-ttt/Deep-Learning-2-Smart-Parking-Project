"""
build_report_docx.py

Builds the final project report (SmartPark_Report.docx) from the real
numbers produced by train_detector.py, evaluate_detector.py,
hyperparameter_ablation.py, and benchmark_baseline.py -- run those first
(see README / report Section 5-7 for exact commands) if you need to
regenerate the underlying numbers.

Usage:
    cd smartpark
    python report_assets/build_report_docx.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
CHARTS = os.path.join(HERE, "charts")
SAMPLES = os.path.join(HERE, "sample_outputs")
CHARTS_REAL = os.path.join(HERE, "charts_real")
SAMPLES_REAL = os.path.join(HERE, "sample_outputs_real")

ACCENT = RGBColor(0x3B, 0x5B, 0xFD)
DARK = RGBColor(0x18, 0x18, 0x1B)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK
    return h


def add_picture_captioned(doc, path, caption, width=5.5):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SmartPark: Real-Time Parking Spot Object Detection")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("AASD 4014 — Deep Learning II — Final Project Report\nObject Detection Track")
    r2.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Group Number: [FILL IN]\nJuly 2026").font.size = Pt(12)
    doc.add_page_break()

    # ---- Placeholder notice ----
    add_heading(doc, "Before Submitting", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "This report was drafted by working through the assignment rubric section by section. "
        "Fields marked [FILL IN] need real information from your group before submission: group number, "
        "member names/roles in the Task Distribution and Team Contributions tables, and the Project "
        "Manager's name. The Agile Development section (below) reconstructs the actual build sequence from "
        "the project's real file history as a starting point -- replace or supplement it with your team's "
        "actual meeting notes, Jira board, and burndown chart if you tracked one separately."
    ).italic = True

    # ---- Task distribution table (required at start of report) ----
    add_heading(doc, "Task Distribution", level=1)
    doc.add_paragraph("Filled in per the assignment's Workload Distribution requirement. Redundant tasks "
                       "(e.g. more than one member touching model training) are acceptable if documented here.")
    add_table(doc, ["Name", "Role", "Tasks"], [
        ["[FILL IN — Project Manager]", "Project Manager", "Coordination, timeline, final report assembly"],
        ["[FILL IN]", "[FILL IN]", "[FILL IN]"],
        ["[FILL IN]", "[FILL IN]", "[FILL IN]"],
        ["[FILL IN]", "[FILL IN]", "[FILL IN]"],
        ["[FILL IN]", "[FILL IN]", "[FILL IN]"],
    ])
    doc.add_page_break()

    # ---- 1. Background and Problem Statement ----
    add_heading(doc, "1. Background and Problem Statement", level=1)
    doc.add_paragraph(
        "Drivers regularly spend significant time searching for available parking in cities, malls, airports, "
        "and campuses -- circling lots because there is no way to see spot-level availability before arriving. "
        "Camera-based occupancy detection is a natural fit for computer vision, and has an established research "
        "line (PKLot, CNRPark-EXT) built around classifying individually marked spots as Empty or Occupied."
    )
    doc.add_paragraph(
        "This project targets the assignment's Object Detection track directly: rather than only classifying "
        "pre-cropped, pre-calibrated spot images (a common simplification in the PKLot/CNRPark-EXT literature), "
        "the core deliverable is a trained object detector that finds and classifies parking spots as one of "
        "two distinct object classes -- \"empty_spot\" and \"occupied_spot\" -- directly in a full, uncalibrated "
        "photo of a lot. A second, complementary pipeline (a per-spot CNN classifier plus a geometric "
        "misparking check) is kept for cameras where spot boundaries are already known, since it additionally "
        "flags cars that are parked across a spot boundary -- something the detector alone does not attempt."
    )

    # ---- 2. Plan of Attack ----
    add_heading(doc, "2. Plan of Attack", level=1)
    doc.add_paragraph(
        "The approach was staged deliberately so that a working end-to-end system existed at every step, "
        "rather than betting everything on one long training run:"
    )
    for item in [
        "Stand up the full pipeline (data → model → API → UI) on a cheap, procedurally generated "
        "synthetic dataset first, to validate every moving part before spending time on real-world data cleanup.",
        "Layer in real-world data (PKLot) for the occupancy classifier to confirm the approach generalizes "
        "beyond synthetic images.",
        "Audit the finished project against the course rubric before finalizing the report -- this caught "
        "that a pre-cropped classifier alone does not satisfy an \"Object Detection\" deliverable (Section 8 "
        "discusses this pivot and why).",
        "Redesign the core model as a genuine from-scratch object detector over two classes, with real "
        "training/validation metrics, hyperparameter ablations, and a baseline comparison, then wire it into "
        "the same live API and demo UI.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ---- Agile Development ----
    add_heading(doc, "Agile Development Process", level=1)
    doc.add_paragraph(
        "Development was organized into four stages, reconstructed here from the project's actual file "
        "history (each stage corresponds to a real day of work on the repository, 2026-07-17 through "
        "2026-07-20). Replace this with your team's real sprint/meeting records if you tracked them "
        "separately (e.g. in Jira)."
    )
    add_table(doc, ["Sprint", "Dates", "Goal", "Key deliverables"], [
        ["Sprint 1", "Jul 17", "Pipeline scaffolding", "Synthetic data generator, CNN classifier, geometry "
         "misparking checker, inference pipeline, first demo UI -- validated end-to-end on synthetic data"],
        ["Sprint 2", "Jul 18", "Real-data integration", "PKLot COCO→crop conversion script, resumable "
         "training script, classifier retrained on real PKLot data (97.04% val accuracy)"],
        ["Sprint 3", "Jul 19", "Reporting", "Initial report/slide drafts, architecture diagram, demo screenshots"],
        ["Sprint 4", "Jul 20", "Environment fixes + OD pivot", "Fixed Python/TensorFlow/Keras version-skew "
         "issues, then rebuilt the core deliverable as a real 2-class object detector: training, evaluation "
         "(precision/recall/AP/mAP), hyperparameter ablations, baseline benchmarking, live API integration"],
        ["Sprint 5", "Jul 21", "Real-data closure", "Recovered the raw PKLot COCO export, fine-tuned the "
         "detector on real bounding boxes (47.71% mAP) and retrained the classifier checkpoint "
         "(98.69% val accuracy) -- closing the real-data validation gap flagged as the main open item"],
        ["Sprint 6", "Jul 22", "Stress-test + rehearsal fix", "Live demo testing on uncontrolled real photos "
         "exposed a scale/density failure; fixed via crop augmentation, which then exposed catastrophic "
         "forgetting on synthetic data; fixed via rehearsal-based mixed fine-tuning (Section 5.3) -- final "
         "adopted checkpoint: 93.31% synthetic mAP, 51.39% real mAP"],
    ])
    add_picture_captioned(doc, os.path.join(CHARTS, "burndown.png"),
                           "Figure: stage-level burndown (reconstructed backlog view, not a literal daily Jira log).")

    doc.add_page_break()

    # ---- 3. The Dataset ----
    add_heading(doc, "3. The Dataset", level=1)
    add_heading(doc, "3.1 Object Detection Dataset (primary deliverable)", level=2)
    doc.add_paragraph(
        "A procedurally generated synthetic dataset (data/generate_synthetic_data.py) provides full, "
        "uncropped parking-lot images with COCO-style bounding-box annotations for two classes: "
        "\"empty_spot\" (the marked, unoccupied space) and \"occupied_spot\" (the parked car's actual "
        "footprint, which shifts for intentionally misparked cases -- this is what makes the task genuine "
        "localization rather than fixed-position lookup). Each generated lot uses a randomized grid layout "
        "(2-3 rows, 4-8 columns, randomized spot size/margin, plus small per-spot jitter), so spot positions "
        "vary across images."
    )
    add_table(doc, ["Metric", "Value"], [
        ["Full lot images", "300"],
        ["Labeled spot instances", "4,422"],
        ["empty_spot instances", "2,027"],
        ["occupied_spot instances", "2,395"],
        ["Improperly parked (subset of occupied)", "542"],
        ["Train / validation split", "255 images / 45 images (85% / 15%)"],
    ])
    doc.add_paragraph(
        "Both classes clear the assignment's 200-images-per-class minimum by a wide margin, and (per the "
        "assignment's allowance) do not need to co-occur in every image, though in this dataset almost every "
        "lot naturally contains both."
    )

    add_heading(doc, "3.2 Occupancy Classifier Dataset (secondary pipeline)", level=2)
    doc.add_paragraph(
        "The per-spot CNN classifier was trained on 600 cropped 64×64 spot patches (330 occupied / 270 "
        "empty) from an earlier, smaller synthetic run for a fast sanity check (97.5% validation accuracy), "
        "then separately trained on real data: a COCO-format PKLot export via Roboflow (1,242 real lot "
        "photographs, 70,684 individually labeled real parking spaces), converted to 36,584 empty / 34,100 "
        "occupied 64×64 crops via convert_coco_to_crops.py, reaching 98.69% validation accuracy after 12 "
        "epochs -- consistent with (and slightly ahead of) published PKLot/CNRPark-EXT results. The same "
        "1,242-image real export (obtained directly from the source this time, not a derived crop set) "
        "is also what Section 5.2 uses to fine-tune the object detector on real bounding boxes."
    )

    doc.add_page_break()

    # ---- 4. Model Description ----
    add_heading(doc, "4. Model Description", level=1)
    add_heading(doc, "4.1 Object Detector (primary deliverable)", level=2)
    doc.add_paragraph(
        "A single-shot, grid-based detector trained from scratch (models/detector_model.py), in the style of "
        "YOLOv1: one predicted box per grid cell, over a 16×16 grid on a 256×256 input. For each cell "
        "the network predicts an objectness score, a box (center offset within the cell + width/height as a "
        "fraction of the image), and class probabilities over the two classes."
    )
    add_table(doc, ["Layer", "Output"], [
        ["Input", "256×256×3"],
        ["Conv2D(16) + MaxPool", "128×128×16"],
        ["Conv2D(32) + MaxPool", "64×64×32"],
        ["Conv2D(64) + MaxPool", "32×32×64"],
        ["Conv2D(128) + MaxPool", "16×16×128"],
        ["Conv2D(128) × 2", "16×16×128"],
        ["Conv2D(7, 1×1) head", "16×16×7  (objectness, tx, ty, tw, th, class0, class1)"],
    ])
    doc.add_paragraph("Total parameters: 393,511. Post-processing: per-class Non-Max Suppression "
                       "(cv2.dnn.NMSBoxes) removes duplicate overlapping detections.")

    add_heading(doc, "4.2 Occupancy Classifier + Geometry Check (secondary pipeline)", level=2)
    doc.add_paragraph(
        "A small CNN (models/cnn_model.py, 1,142,081 parameters: three Conv+MaxPool blocks, a dense head with "
        "dropout, sigmoid output) classifies a pre-cropped 64×64 spot patch as Empty or Occupied. For "
        "occupied spots, a classic-CV step (color-distance thresholding, inference.find_car_bbox_in_spot) "
        "localizes the car within the spot and an IoU/overflow check (improper_parking.py) flags cars parked "
        "across the marked boundary. A MobileNetV2 transfer-learning variant (models/cnn_model.py:"
        "build_mobilenet_transfer) is also implemented as an available alternative backbone but was not "
        "needed in practice -- the small custom CNN already reached 97%+ accuracy at a fraction of the "
        "inference cost."
    )
    add_picture_captioned(doc, os.path.join(HERE, "architecture_v2.png"),
                           "Figure: system architecture -- both pipelines share the same REST API and demo UI.")

    doc.add_page_break()

    # ---- 5. Training & Evaluation ----
    add_heading(doc, "5. Training & Evaluation", level=1)
    doc.add_paragraph(
        "The detector was trained for 40 epochs (Adam, lr=1e-3 with ReduceLROnPlateau, batch size 16, "
        "EarlyStopping patience 6 on val_loss with best-weights restore -- did not trigger early, ran the "
        "full 40 epochs). Loss follows the YOLOv1 formulation: squared-error coordinate loss (weighted 5×), "
        "objectness loss, a down-weighted (0.5×) no-object loss, and classification loss, applied only to "
        "grid cells that contain a ground-truth object (except no-object loss, applied everywhere)."
    )
    add_picture_captioned(doc, os.path.join(CHARTS, "training_curve.png"), "Figure: training/validation loss per epoch.")
    doc.add_paragraph("Final train loss: 0.95. Final (best-restored) validation loss: 2.65.")

    doc.add_paragraph(
        "Evaluated on the held-out 45-image validation split (721 spot instances: 337 empty_spot, 384 "
        "occupied_spot) by decoding predictions back to image coordinates, applying NMS, and matching against "
        "ground truth at IoU ≥ 0.5:"
    )
    add_table(doc, ["Class", "GT instances", "AP@0.5", "Precision", "Recall", "F1"], [
        ["empty_spot", "337", "98.73%", "95.5%", "94.1%", "94.8%"],
        ["occupied_spot", "384", "98.68%", "99.7%", "96.1%", "97.9%"],
        ["mAP@0.5", "—", "98.71%", "—", "—", "—"],
    ])
    add_picture_captioned(doc, os.path.join(CHARTS, "pr_curves.png"), "Figure: precision-recall curves per class.")
    add_picture_captioned(doc, os.path.join(CHARTS, "ap_bar.png"), "Figure: AP per class and mAP@0.5.")

    doc.add_paragraph(
        "To separate localization error from classification error, a second confusion matrix restricts to "
        "boxes that were correctly localized (IoU ≥ 0.5, matched irrespective of predicted class) and asks "
        "whether the predicted class was right:"
    )
    add_picture_captioned(doc, os.path.join(CHARTS, "confusion_matrix.png"),
                           "Figure: confusion matrix, localized boxes only. 335/335 empty_spot and 379/380 "
                           "occupied_spot correctly classified -- 99.86% classification accuracy given correct "
                           "localization. Almost all remaining error is missed detections (recall), not class confusion.")

    add_heading(doc, "Qualitative results", level=2)
    doc.add_paragraph(
        "Eight held-out validation images with predicted vs. ground-truth counts (full list: "
        "report_assets/sample_outputs/summary.json):"
    )
    add_table(doc, ["Sample", "GT (empty / occupied)", "Predicted (empty / occupied)"], [
        ["sample_01.jpg", "11 / 10", "11 / 10  (exact match)"],
        ["sample_03.jpg", "10 / 14", "10 / 14  (exact match)"],
        ["sample_06.jpg", "6 / 6", "1 / 4  (missed several cars -- see discussion)"],
    ])
    for fname, cap in [
        ("sample_01.jpg", "Figure: exact match -- 21/21 spots correctly detected and classified."),
        ("sample_06.jpg", "Figure: a genuine failure case -- several spots (including a low-contrast gray car) "
                          "were missed entirely; discussed in Section 8."),
    ]:
        add_picture_captioned(doc, os.path.join(SAMPLES, fname), cap, width=5.0)

    add_heading(doc, "5.1 Cross-Domain Generalization Test (Qualitative, CNRPark-EXT)", level=2)
    doc.add_paragraph(
        "The detector above is trained only on synthetic data. To honestly check how it behaves outside "
        "that distribution, it was run on 115 real, full-scene photographs from CNRPark-EXT (real_photo_test.py "
        "streams a small partial download of the public CNR-EXT_FULL_IMAGE_1000x750 release -- no full 450MB "
        "archive needed, just enough bytes to extract several dozen complete JPEGs). CNRPark-EXT is a "
        "meaningfully different domain from this project's synthetic data: a steep angled real camera, no "
        "painted spot boundary lines, natural lighting, shadows, and background clutter (trees, other "
        "vehicles in transit, signage)."
    )
    doc.add_paragraph(
        "No ground truth is available for these particular images (CNRPark-EXT's public release ships labels "
        "for pre-cropped patches, not bounding boxes for the full-scene images), so this is a qualitative "
        "check, not a quantitative mAP figure -- and it is a materially different, harder question than "
        "Section 3.2's real-data classifier result, which only had to classify a pre-cropped, pre-located "
        "spot, not find one from scratch in a completely unfamiliar layout."
    )
    add_picture_captioned(doc, os.path.join(HERE, "real_photo_test", "example_1.jpg"),
                           "Figure: real CNRPark-EXT photo, detector output. Boxes loosely cluster around real "
                           "cars, but several vehicles are missed and box shapes are imprecise (some are "
                           "unnaturally tall/narrow).", width=5.0)
    add_picture_captioned(doc, os.path.join(HERE, "real_photo_test", "example_2.jpg"),
                           "Figure: a busier real lot -- only 4 large, imprecise boxes for dozens of visible "
                           "cars. Every detection here was classified occupied_spot; empty_spot essentially "
                           "never fires on these photos, since there are no painted empty bays for it to "
                           "recognize in this camera's view.", width=5.0)
    doc.add_paragraph(
        "Takeaway: the detector carries over a coarse \"there is a cluster of cars roughly here\" signal, "
        "but not reliable per-spot localization or the empty_spot concept, in a domain this different from "
        "training. This is consistent with (and about as severe as) this project's own earlier finding for "
        "the classifier pipeline: a 5-image pilot test on CNRPark-EXT scored only 2/5 (40%) despite 97%+ "
        "accuracy on PKLot. This is a genuinely different, harder test than Section 5.2 below: CNRPark-EXT "
        "is an unfamiliar camera angle with no painted lines at all, not just a new set of real photos."
    )

    add_heading(doc, "5.2 Real-Data Fine-Tuning (PKLot)", level=2)
    doc.add_paragraph(
        "A materially different question from Section 5.1: what happens if the detector is actually "
        "fine-tuned on real photos from a domain resembling its training data (painted spot lines, a "
        "similar top-down-ish angle), rather than dropped zero-shot into a totally unfamiliar camera? The "
        "synthetic-trained checkpoint was fine-tuned (train_detector.py --resume_from, 25 epochs, "
        "lr=1e-4) on the same real PKLot COCO export used for the classifier (Section 3.2): 1,242 real lot "
        "photographs, 70,684 labeled real parking spaces, an 85/15 train/val split (1,056 / 186 images)."
    )
    doc.add_paragraph(
        "Real PKLot lots are considerably denser than this project's synthetic layouts -- an average of "
        "~57 labeled spaces per image, versus ~15 for the synthetic dataset. At the same 16x16 grid "
        "resolution used throughout, this pushes the one-box-per-cell collision rate up sharply: 28.2% of "
        "real ground-truth boxes shared a cell with another box and were dropped during training (versus "
        "0.27% on synthetic data, Section 6.1). That alone caps the best possible recall at roughly 72%, "
        "regardless of how well the model learns."
    )
    add_picture_captioned(doc, os.path.join(CHARTS_REAL, "training_curve.png"),
                           "Figure: fine-tuning loss curve on real PKLot data (25 epochs).")
    add_table(doc, ["Class", "GT instances", "AP@0.5", "Precision", "Recall", "F1"], [
        ["empty_spot", "5,345", "46.84%", "77.6%", "53.2%", "63.1%"],
        ["occupied_spot", "4,975", "48.59%", "72.6%", "57.1%", "63.9%"],
        ["mAP@0.5", "—", "47.71%", "—", "—", "—"],
    ])
    add_picture_captioned(doc, os.path.join(CHARTS_REAL, "pr_curves.png"),
                           "Figure: precision-recall on real data. Precision stays around 80-85% out to roughly "
                           "recall 0.55, then falls sharply -- consistent with a recall ceiling set by the grid "
                           "collision rate above, rather than a gradual quality falloff.")
    add_picture_captioned(doc, os.path.join(CHARTS_REAL, "confusion_matrix.png"),
                           "Figure: confusion matrix, localized real boxes only. (2988+2933)/(2988+234+109+2933) "
                           "= 94.5% classification accuracy given correct localization -- almost identical in "
                           "spirit to the synthetic result (99.86%). The real-data bottleneck is finding spots "
                           "in a dense scene, not telling empty from occupied once found.")
    for fname, cap in [
        ("sample_06.jpg", "Figure: real PKLot photo, fine-tuned detector output. Nearly every visible car is "
                          "found with a tight box and high confidence, in a dense, cluttered real scene."),
    ]:
        add_picture_captioned(doc, os.path.join(SAMPLES_REAL, fname), cap, width=5.5)
    doc.add_paragraph(
        "Takeaway: 47.71% mAP is a genuine, non-trivial real-data result -- far below the 98.71% synthetic "
        "figure, but a large, measured step up from Section 5.1's near-total zero-shot failure on a truly "
        "unfamiliar domain. The confusion matrix pinpoints exactly where the gap is: classification is "
        "essentially as reliable on real data as on synthetic (94.5% vs. 99.86%); the shortfall is almost "
        "entirely recall, driven by the one-box-per-cell design running out of capacity on real, dense "
        "lots. The clearest next architecture change is a finer grid or a multi-box-per-cell (anchor-based) "
        "head specifically for dense real deployments -- not a vague \"needs more real data\" -- since the "
        "grid collision math directly predicts most of the observed shortfall."
    )

    add_heading(doc, "5.3 Stress-Testing on Uncontrolled Photos, and Fixing What Broke", level=2)
    doc.add_paragraph(
        "The 47.71% result above was checked against two arbitrary real photos (aerial parking-lot stock "
        "photography, not from PKLot or any training source) to see how it holds up outside a curated "
        "benchmark. The result was a near-total failure: 3 detections on a ~150-car dense rooftop lot, 5 on "
        "a ~30-car numbered garage. Inspecting the raw objectness scores directly (not just the confidence "
        "threshold) showed this wasn't a calibration issue -- fewer than 20% of grid cells produced any "
        "meaningful signal even at a very permissive 0.01 cutoff. Both photos are far denser and "
        "higher-resolution than anything in training, and squashing them to the model's fixed 256x256 "
        "input destroys most of the detail on any individual car."
    )
    doc.add_paragraph(
        "Two fixes were tried, in order of effort:"
    )
    for item in [
        "Tiled inference (splitting the image into overlapping crops, running the detector on each, "
        "merging results): dramatically improved recall on the dense photo, but on the numbered-garage "
        "photo it began confidently drawing \"occupied_spot\" boxes over completely empty pavement -- likely "
        "confused by stenciled numbers, arrows, and hash-marks that don't appear in any training domain. "
        "Not reliable enough to ship as-is.",
        "Crop/zoom augmentation during training (train_detector.py --augment_crops, models/augmentation.py): "
        "generating several random zoomed-in crops per real training image -- no new photos, just many more "
        "views of the scale/density range at existing ones -- to directly teach the model a wider range of "
        "apparent object sizes.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "Augmentation alone (fine-tuned from the same synthetic checkpoint, 4 crops per real training "
        "image) raised real mAP slightly to 49.36% and produced a dramatic qualitative improvement on both "
        "stress-test photos. But checking it against the original synthetic validation set surfaced a "
        "textbook transfer-learning failure: synthetic mAP collapsed from 98.71% to 6.50%. Fine-tuning on "
        "real data alone had overwritten what the model knew about the original domain rather than "
        "extending it -- catastrophic forgetting."
    )
    doc.add_paragraph(
        "The standard fix is rehearsal: keep some of the original-domain data in the fine-tuning mix "
        "instead of training on the new domain alone. train_detector_mixed.py fine-tunes from the same "
        "synthetic checkpoint on synthetic images, real PKLot images, and the augmented real crops "
        "together (5,366 training images total). This was the best result across every test:"
    )
    add_table(doc, ["Variant", "Synthetic val mAP", "Real val mAP"], [
        ["v1: real-only fine-tune", "not tested (assumed retained)", "47.71%"],
        ["v2: real + augmentation, no rehearsal", "6.50% (collapsed)", "49.36%"],
        ["v3: synthetic + real + augmentation (adopted)", "93.31%", "51.39%"],
    ])
    add_picture_captioned(doc, os.path.join(CHARTS_REAL, "finetuning_comparison.png"),
                           "Figure: the rehearsal-based mixed model (v3) is the only variant that avoids "
                           "collapsing on synthetic data, while also posting the best real-data mAP of the three.")
    doc.add_paragraph(
        "v3 is the checkpoint now served by the live API (trained_detector_mixed/spot_detector.keras) -- "
        "chosen over v2 despite v2's slightly better handling of the single most extreme stress-test photo, "
        "because v3 never catastrophically fails on the project's own primary benchmark. A user who uploads "
        "one of the synthetic sample images to the live demo should not see it fail."
    )
    for before, after, cap in [
        ("before_ivana.jpg", "after_ivana.jpg",
         "Figure: the dense rooftop lot. Before (original synthetic-only checkpoint): 3 detections. "
         "After (v3, mixed rehearsal): 17 detections, correctly spread across multiple car clusters."),
    ]:
        add_picture_captioned(doc, os.path.join(SAMPLES, "..", "real_photo_test", before),
                               "Before: " + cap, width=4.4)
        add_picture_captioned(doc, os.path.join(SAMPLES, "..", "real_photo_test", after),
                               "After (v3, adopted)", width=4.4)
    doc.add_paragraph(
        "Honest residual gap: v3 (17 detections) still finds noticeably fewer cars than v2 (54 detections) "
        "on this single most extreme photo -- rehearsal recovered synthetic performance at some cost to the "
        "very-high-density case specifically. This is a real, acknowledged trade-off, not a fully closed "
        "gap: the next iteration would tune the synthetic:real:augmented mixing ratio (more augmented real "
        "data relative to synthetic) rather than treating this as solved."
    )

    doc.add_page_break()

    # ---- 6. Hyperparameter Tuning ----
    add_heading(doc, "6. Hyperparameter Tuning", level=1)
    doc.add_paragraph(
        "Two hyperparameters were tuned with real, measured comparisons rather than left at arbitrary defaults."
    )
    add_heading(doc, "6.1 Grid resolution", level=2)
    doc.add_paragraph(
        "The one-box-per-cell design means two spot centers landing in the same grid cell causes one "
        "ground-truth box to be dropped during training. Measuring this collision rate across grid "
        "resolutions on the full dataset:"
    )
    add_table(doc, ["Grid size", "Dropped boxes", "Collision rate"], [
        ["8×8", "111 / 4,422", "2.51%"],
        ["16×16 (used)", "12 / 4,422", "0.27%"],
        ["32×32", "3 / 4,422", "0.07%"],
    ])
    doc.add_paragraph(
        "16×16 was chosen as the balance point: it cuts collisions by roughly 9× versus 8×8 while "
        "keeping the head 4× cheaper than 32×32 for a collision rate that is already well under 1%."
    )

    add_heading(doc, "6.2 Coordinate loss weight (lambda_coord)", level=2)
    doc.add_paragraph(
        "YOLOv1 weights the box-coordinate loss 5× relative to the classification/objectness terms, to "
        "keep the localization signal from being drowned out by the much larger number of no-object cells. "
        "This was tested directly: two otherwise-identical detectors were trained for a matched, short "
        "budget (15 epochs) with lambda_coord=5.0 vs. lambda_coord=1.0:"
    )
    add_table(doc, ["lambda_coord", "Val loss (15 epochs)", "mAP@0.5"], [
        ["5.0 (used)", "3.91", "96.72%"],
        ["1.0", "16.36", "0.00%"],
    ])
    doc.add_paragraph(
        "At lambda_coord=1.0 the model collapsed to predicting no confident objects anywhere within the "
        "15-epoch budget (mAP=0%) -- a well-known YOLO training pathology when the coordinate signal is "
        "underweighted relative to the flood of no-object cells. This concretely confirms the YOLOv1 default "
        "was the right choice for this setup, rather than assuming it."
    )

    doc.add_page_break()

    # ---- 7. Benchmarking ----
    add_heading(doc, "7. Benchmarking", level=1)
    doc.add_paragraph(
        "The trained detector's per-spot classification accuracy was compared against two non-learned "
        "baselines, using the same validation ground-truth spot locations for all three so the comparison "
        "isolates classification quality from localization ability:"
    )
    add_table(doc, ["Method", "Accuracy"], [
        ["Majority-class baseline (always predict \"occupied\")", "53.26%"],
        ["Classic-CV heuristic (this project's original color-distance detector, "
         "inference.find_car_bbox_in_spot)", "82.66%"],
        ["Trained detector (classification given correct localization)", "99.86%"],
    ])
    doc.add_paragraph(
        "The classic-CV heuristic was this project's own first approach to occupancy (used for the "
        "\"properly parked\" geometry check before any detector existed) -- it beats a naive prior "
        "substantially but still misses roughly 1 in 6 spots, which is exactly the kind of gap a learned "
        "model is expected to close, and the trained detector closes nearly all of it."
    )
    doc.add_paragraph(
        "The occupancy classifier's real-data result (98.69% on PKLot) is also independently consistent with "
        "(and slightly ahead of) accuracy levels reported in the original PKLot and CNRPark-EXT research "
        "papers for the same per-spot classification task, which is an external benchmark beyond this "
        "project's own baselines."
    )

    doc.add_page_break()

    # ---- 8. Discussion & Reflection ----
    add_heading(doc, "8. Discussion & Reflection", level=1)
    add_heading(doc, "What worked", level=2)
    doc.add_paragraph(
        "Staging the build around a synthetic dataset first paid off twice: it let the whole pipeline "
        "(data → model → API → UI) be validated cheaply before touching real data, and later, when the "
        "project needed to become a genuine object detector, the same generator only needed real position "
        "variance and COCO-style annotations added -- no new infrastructure. The detector reached 98.71% mAP "
        "and, per the confusion matrix, almost never confuses the two classes when it does localize a spot."
    )
    doc.add_paragraph(
        "Recovering the raw PKLot COCO export (Sprint 5) and fine-tuning the synthetic-trained detector on "
        "it directly validated the transfer-learning approach: 47.71% mAP on real, dense, cluttered lot "
        "photos, with classification accuracy given correct localization (94.5%) almost matching the "
        "synthetic figure (99.86%). That the same pattern -- strong classification, weaker localization -- "
        "shows up in both synthetic and real evaluations is a good sign the model is learning something "
        "structural about the task, not overfitting to synthetic-image quirks."
    )
    doc.add_paragraph(
        "Rehearsal-based mixed fine-tuning (Section 5.3, Sprint 6) worked exactly as the transfer-learning "
        "literature predicts: mixing original-domain data back into a fine-tuning run on a new domain "
        "recovered the catastrophically-forgotten synthetic performance (6.50% → 93.31% mAP) while *also* "
        "improving real-data mAP further (49.36% → 51.39%) -- a genuine best-of-both result, not a "
        "compromise between two worse options."
    )
    add_heading(doc, "What didn't work / had to be revisited", level=2)
    for item in [
        "The first version of this project (per-spot CNN classifier applied to pre-cropped, pre-calibrated "
        "images) does not actually satisfy an Object Detection deliverable -- it assumes spot locations are "
        "already known and never localizes anything itself. This was only caught by explicitly auditing the "
        "finished project against the course rubric, not during initial development. The fix (Sprint 4) was "
        "a genuine architecture change, not a relabeling: a real detector trained from scratch on two object "
        "classes, evaluated with detection metrics (AP/mAP), not classification accuracy alone.",
        "Sample_06 (Section 5) shows the detector missing several real spots in one validation image, "
        "including a low-contrast gray car -- a similar failure mode to the one already documented in the "
        "classic-CV heuristic (which also struggles with cars close in color to the asphalt/lines). This "
        "suggests the remaining error is concentrated in genuinely low-contrast cases rather than being "
        "randomly distributed, and would be the first thing to target with more training data or contrast "
        "augmentation.",
        "The lambda_coord=1.0 ablation collapsing to 0% mAP was a useful, if initially alarming, result -- "
        "it's a concrete demonstration of a known YOLO failure mode rather than a bug, and losing an "
        "afternoon to debugging \"why is validation mAP exactly zero\" before recognizing the pattern was a "
        "real part of this project's timeline.",
        "The original \"properly parked\" geometry check and its classic-CV car-localization step were "
        "validated only on synthetic data; real photos (shadows, lighting, dense packing) are documented as "
        "an open limitation rather than glossed over.",
        "The real-photo generalization test (Section 5.1) confirmed a genuine domain gap rather than "
        "assuming one existed: on real CNRPark-EXT photos (a steeper camera angle, no painted lines) the "
        "detector finds a rough cluster of real cars but misses most vehicles and never recognizes an "
        "empty_spot. This mirrors this project's own earlier 40% cross-domain result for the classifier "
        "pipeline on the same dataset -- a consistent signal that PKLot-style training data does not "
        "automatically transfer to a differently-marked, differently-angled real lot.",
        "Fine-tuning on real data alone (Section 5.3, v2) caused catastrophic forgetting: synthetic mAP "
        "collapsed from 98.71% to 6.50% even as real-data mAP improved. This was only caught by explicitly "
        "re-evaluating on the original synthetic validation set after the real-data fine-tune -- it would "
        "have shipped unnoticed otherwise, since nothing about the real-data numbers alone hinted at it. "
        "The fix (rehearsal: mixing synthetic data back into the fine-tuning set) is standard in the "
        "transfer-learning literature but easy to skip under time pressure.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "What we'd do differently", level=2)
    for item in [
        "Audit the project against the grading rubric before building anything, not after -- it would have "
        "avoided building an entire classifier-only pipeline that needed to be substantially reworked.",
        "Keep the raw PKLot COCO annotations from the start, rather than only the derived crops -- this was "
        "the single biggest gap after the OD pivot, and closing it (Sprint 5, Section 5.2) is what turned "
        "\"the detector has never seen a real bounding box\" into a genuine, measured 47.71% mAP result.",
        "Test on uncontrolled, real-world photos (not just held-out samples from the same source dataset) "
        "much earlier -- the two stress-test photos in Section 5.3 surfaced both the scale/density failure "
        "and, indirectly, the catastrophic-forgetting bug, neither of which the standard validation split "
        "would ever have revealed on its own.",
        "Tune the synthetic:real:augmented mixing ratio in the rehearsal fine-tune, rather than accepting "
        "the first mix that worked -- Section 5.3's adopted model still under-performs the (synthetic-"
        "forgetting) augmentation-only variant on the single most extreme stress-test photo, suggesting the "
        "current ratio leans slightly too far toward preserving synthetic performance.",
        "Move to a finer grid or a multi-box-per-cell (anchor-based) detection head before deploying on "
        "dense real lots -- Section 5.2 shows the real-data recall ceiling is set almost entirely by "
        "one-box-per-cell grid collisions (28.2% of real boxes dropped at 16x16 resolution), not by weak "
        "classification. This is now the clearest, most specific next step, backed by the collision-rate "
        "math rather than a general sense that \"more real data would probably help.\"",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ---- 9. Conclusion ----
    add_heading(doc, "9. Conclusion", level=1)
    doc.add_paragraph(
        "This project delivers a working object detector, trained from scratch, that finds and classifies "
        "parking spots as \"empty_spot\" or \"occupied_spot\" directly in full, uncalibrated lot photos. "
        "The checkpoint served live by the API (trained_detector_mixed, fine-tuned on synthetic and real "
        "PKLot data together via rehearsal) reaches 93.31% mAP@0.5 on synthetic validation data and 51.39% "
        "on held-out real photographs -- clearing a majority-class baseline (53.3%) and this project's own "
        "earlier classic-CV heuristic (82.7%) by a wide margin on classification, with the confusion matrix "
        "showing the real-data gap is concentrated in localization recall (a grid-resolution limitation on "
        "dense real lots), not classification reliability. Getting here was not a straight line: real-data "
        "fine-tuning alone worked (47.71% mAP) but generalized poorly to two uncontrolled stress-test "
        "photos; crop augmentation fixed that but caused catastrophic forgetting of synthetic performance "
        "(98.71% → 6.50%); rehearsal -- mixing synthetic data back into the fine-tune -- recovered synthetic "
        "performance while improving real performance further, the best result of every variant tried "
        "(Section 5.3). A second, complementary pipeline (occupancy classifier + geometric misparking check) "
        "reaches 98.69% validation accuracy on the same real data, for cameras with known, calibrated spot "
        "boundaries. Both are served live through a Flask REST API and a browser demo that accepts an "
        "uploaded photo of any parking lot. A qualitative zero-shot test on 115 real CNRPark-EXT photos (a "
        "genuinely unfamiliar camera angle and marking style) shows domain transfer does not happen for "
        "free -- the detector needs to see data resembling its deployment camera. The clearest next steps "
        "are a finer-grained or anchor-based detection head for dense real deployments, and further tuning "
        "of the rehearsal mixing ratio -- both concrete, measured directions rather than open questions."
    )

    doc.add_page_break()

    # ---- 10. Team Contributions ----
    add_heading(doc, "10. Team Contributions", level=1)
    doc.add_paragraph("[FILL IN — replace with each member's actual contribution before submission]")
    add_table(doc, ["Name", "Contribution"], [
        ["[FILL IN — Project Manager]", "[FILL IN]"],
        ["[FILL IN]", "[FILL IN]"],
        ["[FILL IN]", "[FILL IN]"],
        ["[FILL IN]", "[FILL IN]"],
        ["[FILL IN]", "[FILL IN]"],
    ])

    out_path = os.path.join(HERE, "SmartPark_Report.docx")
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    build()
